import re
import os
import time
import json
import asyncio
from typing import List
from dotenv import load_dotenv

import win32com.client
from docx import Document
from docx2pdf import convert

from openai import AsyncOpenAI
from fastapi import FastAPI, HTTPException, Query
from fastapi.responses import JSONResponse
import uvicorn

# Load environment variables from .env file for API keys, etc.
load_dotenv()
api_key = os.getenv("API_KEY")

# Initialize OpenAI client for GPT-based grammar correction
client = AsyncOpenAI(api_key=api_key)

# ---------------------------
# PDF to Word Conversion
# ---------------------------
def convert_pdf_to_word(pdf_path, docx_path):
    """
    Convert a PDF file to a Word (.docx) file using Microsoft Word automation.
    """
    print("Converting PDF to Word using Microsoft Word...")
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(pdf_path)
    doc.SaveAs(docx_path, FileFormat=16)
    doc.Close()
    word.Quit()
    print("Conversion done.")

# ---------------------------
# Tokenizer with punctuation
# ---------------------------
def tokenize_with_punctuation(text):
    """
    Tokenizer that splits punctuation as separate tokens.
    """
    return re.findall(r"\w+|[^\w\s]", text, re.UNICODE)

# ---------------------------
# Word to PDF Conversion
# ---------------------------
def convert_word_to_pdf(updated_docx_path, final_pdf_path):
    """
    Convert a Word (.docx) file back to PDF.
    """
    print("Converting back to PDF...")
    convert(updated_docx_path, final_pdf_path)
    print(f"Final PDF saved as: {final_pdf_path}")

# ---------------------------
# Short/uncorrectable text filter
# ---------------------------
def is_short_or_uncorrectable(text):
    """
    Returns True if the text is too short, a URL, a number, or contains no letters.
    Used to skip unnecessary GPT calls.
    """
    tokens = re.findall(r"\w+", text)
    if len(tokens) < 5:
        if text.strip().startswith("http") or re.match(r"^\d+(\.\d+)*$", text.strip()):
            return True
        if not re.search(r"[a-zA-Z]", text):
            return True
    return False

# ---------------------------
# GPT-based grammar correction
# ---------------------------
async def gpt_proofread(text):
    """
    Calls OpenAI GPT to proofread and classify changes in a sentence.
    Returns a structured JSON object with corrections and diffs.
    """
    if is_short_or_uncorrectable(text):
        # Return fallback JSON directly, no call to GPT
        return {
            "original": text,
            "corrected": text,
            "original_token": [],
            "proofread_token": [],
            "changes": []
        }
    text = re.sub(r"[^\S\r\n]{2,}", " ", text)
    system_msg = (
        "You are a formal grammar corrector and change classifier. "
        "You will receive one sentence at a time and must return a strictly formatted JSON object with the following keys:\n\n"
        "- 'original': the exact original input sentence as a string.\n"
        "- 'corrected': the fully corrected version of the sentence as a string.\n"
        "- 'original_token': a list of token objects from the original sentence, each with:\n"
        "    - 'idx': integer index starting at 0\n"
        "    - 'word': the exact word or punctuation string\n"
        "- 'proofread_token': a list of token objects from the corrected sentence, each with:\n"
        "    - 'idx': integer index starting at 0\n"
        "    - 'word': the exact corrected word or punctuation string\n"
        "- 'changes': a list of grammar or spelling changes detected, where each change is an object containing:\n"
        "    - 'type': one of ['replaced', 'inserted', 'removed', 'corrected']\n"
        "    - 'original_idx': integer index in original_token or null if the change is an insertion\n"
        "    - 'proofread_idx': integer index in proofread_token or null if the change is a removal\n"
        "    - 'original_word': the original word or punctuation involved\n"
        "    - 'proofread_word': the corrected word or punctuation (empty string if removed)\n"
        "    - 'suggestion': up to three synonyms (only for 'replaced' or 'inserted' types; empty list otherwise)\n\n"
        "Important instructions:\n"
        "- Treat punctuation changes (e.g., commas, periods) as valid changes and include them in 'changes'.\n"
        "- Normalize all excessive or irregular spacing by converting multiple spaces into a single space.\n"
        "- Do not include any empty strings or tokens consisting only of whitespace in the token lists.\n"
        "- If the input sentence is very short or contains no detectable errors, return 'corrected' identical to 'original' and empty lists for 'original_token', 'proofread_token', and 'changes'.\n"
        "- You MUST respond with STRICT JSON ONLY. No extra text, no markdown, no explanations.\n"
        "- The JSON must be well-formed and parsable.\n"
        "- If the sentence is very short or cannot be corrected, return exactly:\n"
        "{\n"
        "  \"original\": \"<input>\",\n"
        "  \"corrected\": \"<input>\",\n"
        "  \"original_token\": [tokenized],\n"
        "  \"proofread_token\": [tokenized],\n"
        "  \"changes\": []\n"
        "}\n"
    )
    
    user_msg = f"Original sentence:\n{text}\n\nPlease return only the JSON."

    response = await client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.3,
    )

    content = response.choices[0].message.content.strip()
    # print("GPT raw response:", content)  # Uncomment for debugging

    # Improved JSON extraction and parsing
    try:
        # Try to parse the whole response first
        result = json.loads(content)
        return result
    except Exception:
        try:
            # Try to extract the first JSON object in the response
            match = re.search(r"\{.*\}", content, re.DOTALL)
            if not match:
                raise ValueError("No JSON object found in GPT response")
            json_str = match.group(0)
            result = json.loads(json_str)
            return result
        except Exception as e:
            print(f"Failed to parse GPT response for input: {text}\nError: {e}")
            # Return a fallback result so processing continues
            return {
                "original": text,
                "corrected": text,
                "original_token": [],
                "proofread_token": [],
                "changes": []
            }

# ---------------------------
# Async wrapper for GPT proofreading
# ---------------------------
async def async_gpt_proofread(paragraph_id, text):
    """
    Async wrapper to call GPT proofreading for a paragraph.
    Returns paragraph_id, original text, and GPT response.
    """
    try:
        gpt_response = await gpt_proofread(text)
        return paragraph_id, text, gpt_response
    except Exception as e:
        print(f"Error processing paragraph {paragraph_id}: {e}")
        return paragraph_id, text, {"corrected": text}
    
# ---------------------------
# Main async grammar correction for all paragraphs
# ---------------------------
async def correct_paragraphs_async(
    docx_path, updated_docx_path, json_output_path, pdf_id="example_pdf_001"
):
    """
    Proofreads all paragraphs in a Word document asynchronously using GPT,
    updates the document, and writes a JSON report.
    Returns the number of improved paragraphs.
    """
    doc = Document(docx_path)
    data = {"pdf_id": pdf_id, "paragraphs": []}
    total_word_changes = 0

    tasks = []
    paragraph_map = []

    # Step 1: prepare tasks for each paragraph
    for idx, paragraph in enumerate(doc.paragraphs):
        original_text = paragraph.text
        if not original_text.strip():
            continue
        para_id = idx + 1
        paragraph_map.append((idx, para_id, paragraph))
        tasks.append(async_gpt_proofread(para_id, original_text))

    # Step 2: execute all tasks concurrently
    results = await asyncio.gather(*tasks)

      # Step 3: update paragraphs and prepare report
    for (idx, para_id, paragraph), (returned_para_id, original_text, gpt_response) in zip(paragraph_map, results):
        corrected_text = gpt_response.get("corrected", original_text)

        data["paragraphs"].append(
            {
                "paragraph_id": para_id,
                "original": gpt_response.get("original"),
                "proofread": gpt_response.get("corrected"),
                "original_token": gpt_response.get("original_token", []),
                "proofread_token": gpt_response.get("proofread_token", []),
                "original_text": [
                    {
                        "index": ch.get("original_idx"),
                        "word": ch.get("original_word"),
                        "type": "error",
                    }
                    for ch in gpt_response.get("changes", [])
                    if ch.get("type") != "inserted"
                    and ch.get("original_idx") is not None
                ],
                "revised_text": [
                    {
                        "index": ch.get("proofread_idx"),
                        "word": ch.get("proofread_word"),
                        "type": ch.get("type"),
                        "suggestions": ch.get("suggestion", [ch.get("proofread_word")]),
                    }
                    for ch in gpt_response.get("changes", [])
                    if ch.get("proofread_idx") is not None
                ],
            }
        )

        # Update Word paragraph while preserving formatting
        if paragraph.runs:
            ref_run = paragraph.runs[0]
            paragraph.clear()
            new_run = paragraph.add_run(corrected_text)
            new_run.font.name = ref_run.font.name
            new_run.bold = ref_run.bold
            new_run.italic = ref_run.italic
            new_run.underline = ref_run.underline
            new_run.font.size = ref_run.font.size
            if ref_run.font.color and ref_run.font.color.rgb:
                new_run.font.color.rgb = ref_run.font.color.rgb
        else:
            paragraph.text = corrected_text

        # Count word changes for this paragraph and accumulate
        word_changes_count = sum(
            1 for ch in gpt_response.get("changes", [])
            if ch.get("type") in ["inserted", "changed", "replaced"]
        )
        total_word_changes += word_changes_count

    doc.save(updated_docx_path)
    print("Document updated with grammar corrections.")

    with open(json_output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print(f"Proofread JSON saved to {json_output_path}")

    return total_word_changes

# ---------------------------
# Update selected paragraphs in DOCX from JSON
# ---------------------------
def update_changes_on_pdf(
    final_pdf_path, updated_docx_path, json_output_path, paragraph_id
):
    """
    Updates the paragraphs in a Word document based on proofread data from a JSON file.
    Only updates paragraphs whose IDs are in paragraph_id.
    Returns the number of paragraphs updated.
    """
    # Load proofreading results from JSON
    with open(json_output_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    paragraphs_data = data.get("paragraphs", [])
    paragraph_id_set = set(paragraph_id)

    # Load the existing DOCX file
    doc = Document(updated_docx_path)
    updated_count = 0

    for para in paragraphs_data:
        pid = str(para.get("paragraph_id"))
        if pid in paragraph_id_set:
            print(f"Now Processing Paragraph ID : {pid}")
            proofread_text = para.get("proofread")
            para_index = int(pid) - 1  # Adjust for 0-based indexing

            if 0 <= para_index < len(doc.paragraphs):
                paragraph = doc.paragraphs[para_index]
                if paragraph.runs:
                    ref_run = paragraph.runs[0]
                    paragraph.clear()
                    new_run = paragraph.add_run(proofread_text)
                    new_run.font.name = ref_run.font.name
                    new_run.bold = ref_run.bold
                    new_run.italic = ref_run.italic
                    new_run.underline = ref_run.underline
                    new_run.font.size = ref_run.font.size
                    if ref_run.font.color and ref_run.font.color.rgb:
                        new_run.font.color.rgb = ref_run.font.color.rgb
                else:
                    paragraph.text = proofread_text
                updated_count += 1

    # Save changes to DOCX
    doc.save(updated_docx_path)

    return updated_count

# ---------------------------
# FastAPI app and endpoint
# ---------------------------
app = FastAPI()

@app.get("/api/grammar-check")
async def grammar_check(
    mode: int = Query(...), file_code: str = Query(...), paragraph_id: str = Query(...)
):
    """
    Main API endpoint for grammar checking and PDF processing.
    - mode="0": Full process (PDF→Word→Proofread→PDF)
    - mode="1": Update only selected paragraphs (using paragraph_id)
    Returns output filenames, total improvements, and elapsed time.
    """
    print(
        f"Received request: mode={mode}, file_code={file_code}, paragraph_id={paragraph_id}"
    )

    # Clean input like "[1,2,3]" or "1,2,3"
    cleaned = re.sub(r"[\[\]\s]", "", paragraph_id)

    start_time = time.time()

    pdf_path = os.path.abspath(f"original_pdfs/{file_code}.pdf")
    docx_path = os.path.abspath(f"parsing_words/{file_code}_temp.docx")
    updated_docx_path = os.path.abspath(f"parsing_words/{file_code}_updated.docx")
    final_pdf_path = os.path.abspath(f"processed_pdfs/xxx_{file_code}.pdf")
    json_output_path = os.path.abspath(f"jsons/xxx_{file_code}.json")

    if mode == 0:
        # Full process: convert, proofread, and save all
        convert_pdf_to_word(pdf_path, docx_path)
        total_improvements = await correct_paragraphs_async(
            docx_path, updated_docx_path, json_output_path
        )
        convert_word_to_pdf(updated_docx_path, final_pdf_path)
    else:
        # Only update selected paragraphs
        convert_pdf_to_word(final_pdf_path, docx_path)
        total_improvements = update_changes_on_pdf(
            final_pdf_path, updated_docx_path, json_output_path, paragraph_id
        )
        convert_word_to_pdf(updated_docx_path, final_pdf_path)

    elapsed = time.time() - start_time
    print(f"Total processing time: {elapsed:.2f} seconds")
    print(f"Total Improvements Found: {total_improvements}")

    return {
        "json_filename": os.path.basename(json_output_path),
        "final_pdf_filename": os.path.basename(final_pdf_path),
        "total_improvements": total_improvements,
        "elapsed_time_seconds": round(elapsed, 2),
    }

if __name__ == "__main__":
    # Run the FastAPI app with Uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)